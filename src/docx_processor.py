import docx
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK
from src.kanconverter import process_line  # Import the function from your original script

def convert_paragraph(paragraph):
    """Convert the text in a paragraph from ASCII to Unicode Kannada."""
    converted_runs = []
    for run in paragraph.runs:
        converted_text = process_line(run.text)
        new_run = paragraph.add_run(converted_text)
        new_run.font.name = run.font.name
        new_run.font.size = run.font.size
        new_run.font.bold = run.font.bold
        new_run.font.italic = run.font.italic
        new_run.font.underline = run.font.underline
        converted_runs.append(new_run._element)
    
    # Clear original runs and add converted runs
    for run in paragraph.runs[:-len(converted_runs)]:
        run._element.getparent().remove(run._element)

def convert_table(table):
    """Convert the text in a table from ASCII to Unicode Kannada."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                convert_paragraph(paragraph)

def convert_docx(input_file, output_file):
    """Convert a .docx file from ASCII to Unicode Kannada while preserving structure."""
    doc = docx.Document(input_file)
    
    # Convert main document content
    for element in doc.element.body:
        if isinstance(element, docx.text.paragraph.Paragraph):
            convert_paragraph(element)
        elif isinstance(element, docx.table.Table):
            convert_table(element)
    
    # Convert headers and footers
    for section in doc.sections:
        for header in section.header.paragraphs:
            convert_paragraph(header)
        for footer in section.footer.paragraphs:
            convert_paragraph(footer)

    # Save the converted document
    doc.save(output_file)

# if __name__ == "__main__":
#     input_file = "input.docx"
#     output_file = "output_unicode.docx"
#     convert_docx(input_file, output_file)
#     print(f"Conversion complete. Unicode Kannada document saved as {output_file}")